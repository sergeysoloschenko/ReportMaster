"""
Word Document Generator
Creates formal monthly reports following the consultant/operator section template
"""

import logging
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List


class WordReportGenerator:
    """Generate Word document reports in formal structure"""
    
    # Russian translations
    TRANSLATIONS_RU = {
        'title': 'Ежемесячный отчет',
        'section_title': '4. Работа с консультантами и операторами',
        'generated': 'Создано',
        'context': 'Контекст:',
        'actions': 'Действия:',
        'result': 'Результат / Статус:',
        'period': 'Период / Даты:',
        'parties': 'Стороны / Контрагенты:',
        'remarks': 'Замечания / Риски:',
        'recommendations': 'Рекомендации / Следующие шаги:',
        'statistics': 'Статистика отчета',
        'total_categories': 'Всего подразделов',
        'total_messages': 'Всего сообщений',
        'total_attachments': 'Всего вложений'
    }
    
    # Month names in Russian (genitive case for report title)
    MONTHS_RU = {
        1: 'января', 2: 'февраля', 3: 'марта', 4: 'апреля',
        5: 'мая', 6: 'июня', 7: 'июля', 8: 'августа',
        9: 'сентября', 10: 'октября', 11: 'ноября', 12: 'декабря'
    }
    
    def __init__(self, config: dict):
        self.logger = logging.getLogger(__name__)
        self.config = config
        
        report_config = config.get('report', {})
        self.table_style = report_config.get('table_style', {})
        self.font = self.table_style.get('font', 'Calibri')
        self.font_size = self.table_style.get('font_size', 11)
    
    def generate_report(self, summaries: Dict, output_path: Path, report_month: str = None) -> Path:
        """Generate structured Word document report"""
        self.logger.info(f"Generating structured report: {output_path}")
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = self.font
        font.size = Pt(self.font_size)
        
        # Add header
        self._add_header(doc, report_month)
        
        # Add section 4 with subsections
        self._add_section_4(doc, summaries)
        
        # Add statistics
        self._add_statistics(doc, summaries)
        
        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        self.logger.info(f"✓ Report saved: {output_path}")
        
        return output_path
    
    def _add_header(self, doc, report_month: str = None):
        """Add report header"""
        
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(self.TRANSLATIONS_RU['title'])
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(68, 114, 196)
        
        # Subtitle (month/year)
        if not report_month:
            now = datetime.now()
            month_ru = self.MONTHS_RU[now.month].capitalize()
            report_month = f"{month_ru} {now.year}"
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(report_month)
        run.font.size = Pt(14)
        
        # Generation date
        now = datetime.now()
        date_ru = f"{now.day} {self.MONTHS_RU[now.month]} {now.year} г."
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(f"{self.TRANSLATIONS_RU['generated']}: {date_ru}")
        run.font.size = Pt(10)
        run.italic = True
        
        # Add spacing
        doc.add_paragraph()
    
    def _add_section_4(self, doc, summaries: Dict):
        """Add Section 4: Work with consultants and operators"""
        
        # Section title
        section_title = doc.add_heading(self.TRANSLATIONS_RU['section_title'], level=1)
        section_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph()
        
        # Add subsections for each category
        for idx, (cat_id, summary_data) in enumerate(summaries.items(), 1):
            self._add_subsection(doc, idx, summary_data)
    
    def _add_subsection(self, doc, subsection_num: int, summary_data: Dict):
        """Add a subsection (4.x) following the template structure"""
        
        # Subsection heading: 4.x [Category Name]
        heading_text = f"4.{subsection_num} {summary_data['category_name']}"
        heading = doc.add_heading(heading_text, level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Context
        context_para = doc.add_paragraph()
        context_para.add_run(self.TRANSLATIONS_RU['context']).bold = True
        context_para.add_run(f" {summary_data['context']}")
        
        # Actions
        actions_para = doc.add_paragraph()
        actions_para.add_run(self.TRANSLATIONS_RU['actions']).bold = True
        
        actions_list = summary_data.get('actions', [])
        if actions_list:
            for action in actions_list:
                action_item = doc.add_paragraph(action, style='List Number')
                action_item.paragraph_format.left_indent = Inches(0.5)
        else:
            doc.add_paragraph("Переписка по данному вопросу", style='List Number').paragraph_format.left_indent = Inches(0.5)
        
        # Result / Status
        result_para = doc.add_paragraph()
        result_para.add_run(self.TRANSLATIONS_RU['result']).bold = True
        result_text = summary_data.get('result', 'В процессе')
        result_para.add_run(f" {result_text}")
        
        # Period / Dates
        period_para = doc.add_paragraph()
        period_para.add_run(self.TRANSLATIONS_RU['period']).bold = True
        period_para.add_run(f" {summary_data['date_range']}")
        
        # Parties / Contractors
        parties_para = doc.add_paragraph()
        parties_para.add_run(self.TRANSLATIONS_RU['parties']).bold = True
        parties_text = summary_data.get('parties', ', '.join(summary_data['participants'][:5]))
        parties_para.add_run(f" {parties_text}")
        
        # Remarks / Risks (only if present)
        remarks = summary_data.get('remarks', '').strip()
        if remarks:
            remarks_para = doc.add_paragraph()
            remarks_para.add_run(self.TRANSLATIONS_RU['remarks']).bold = True
            remarks_para.add_run(f" {remarks}")
        
        # Recommendations / Next Steps (only if present)
        recommendations = summary_data.get('recommendations', '').strip()
        if recommendations:
            rec_para = doc.add_paragraph()
            rec_para.add_run(self.TRANSLATIONS_RU['recommendations']).bold = True
            rec_para.add_run(f" {recommendations}")
        
        # Add spacing between subsections
        doc.add_paragraph()
    
    def _add_statistics(self, doc, summaries: Dict):
        """Add report statistics"""
        
        doc.add_page_break()
        
        total_messages = sum(s['message_count'] for s in summaries.values())
        total_attachments = sum(s['attachment_count'] for s in summaries.values())
        
        heading = doc.add_paragraph()
        heading.add_run(self.TRANSLATIONS_RU['statistics']).bold = True

        doc.add_paragraph(f"• {self.TRANSLATIONS_RU['total_categories']}: {len(summaries)}")
        doc.add_paragraph(f"• {self.TRANSLATIONS_RU['total_messages']}: {total_messages}")
        doc.add_paragraph(f"• {self.TRANSLATIONS_RU['total_attachments']}: {total_attachments}")


if __name__ == "__main__":
    print("✓ Word generator module loaded successfully")
